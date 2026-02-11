import streamlit as st
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
import openai
import re
import os
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import time 

# ==========================================
# [설정] 페이지 기본 설정
# ==========================================
st.set_page_config(page_title="사계국어 모의고사 시스템", page_icon="📚", layout="wide") 

# ==========================================
# [설정] API 클라이언트 초기화 (Google + OpenAI 통합)
# ==========================================
try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=GOOGLE_API_KEY)
except (KeyError, AttributeError):
    GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")
    if GOOGLE_API_KEY:
        genai.configure(api_key=GOOGLE_API_KEY) 

openai_client = None
try:
    if "OPENAI_API_KEY" in st.secrets:
        from openai import OpenAI
        openai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception as e:
    pass

# 모델 우선순위 정의
MODEL_PRIORITY = [
    "gpt-5.2"
    "gpt-4o",               
    "gemini-1.5-pro",       
    "gemini-1.5-flash"      
] 

# ==========================================
# [초기화] Session State 설정
# ==========================================
if 'generation_requested' not in st.session_state:
    st.session_state.generation_requested = False 

if 'generated_result' not in st.session_state:
    st.session_state.generated_result = None 

if 'app_mode' not in st.session_state:
    st.session_state.app_mode = "⚡ 비문학 문제 제작" 

# ==========================================
# [공통 HTML/CSS 정의] - 원본 스타일 100% 보존
# ==========================================
HTML_HEAD = """
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <style>
        body { 
            font-family: 'Malgun Gothic', 'Batang', serif; 
            padding: 40px; 
            max-width: 900px; 
            margin: 0 auto; 
            line-height: 1.6; 
            color: #000; 
            font-size: 11pt;
        }
        
        .header-container {
            margin-bottom: 30px;
            border-bottom: 2px solid #000; 
            padding-bottom: 15px;
            text-align: center; 
        }
        
        .top-row {
            display: flex;
            justify-content: space-between;
            align-items: flex-end; 
            margin-bottom: 20px;
        }
        
        .main-title {
            font-size: 26px;
            font-weight: 800;
            margin: 0;
            letter-spacing: -0.5px;
            color: #000;
            line-height: 1.2;
            flex-grow: 1;
            text-align: left; 
        }
        
        .time-box {
            font-size: 14px;
            font-weight: bold;
            border: 1px solid #000;
            padding: 5px 15px;
            border-radius: 4px;
            white-space: nowrap;
        }
        
        .topic-info {
            font-size: 16px;
            font-weight: 800; 
            color: #000;
            background-color: #f4f4f4; 
            padding: 8px 20px;
            display: inline-block;
            border-radius: 8px;
            margin-top: 5px;
        }

        .passage { 
            font-size: 10.5pt; border: 1px solid #444; padding: 30px; 
            margin-bottom: 40px; background-color: #fff; 
            line-height: 1.8; text-align: justify;
        }
        .passage p { text-indent: 0.7em; margin-bottom: 15px; }

        .poetry-passage { 
            white-space: pre-wrap; font-family: 'Batang', serif; line-height: 2.2;
            font-size: 11pt; border: 1px solid #444; padding: 35px;
            margin-bottom: 40px; background-color: #fff; text-align: left;
        }

        /* 분석 차트 스타일 */
        .analysis-chart { width: 100%; border-collapse: collapse; margin-bottom: 40px; table-layout: fixed; }
        .analysis-chart th { 
            background-color: #f8f9fa; border: 1px solid #444; padding: 12px; 
            font-weight: bold; width: 120px; text-align: center; font-size: 10.5pt; 
        }
        .analysis-chart td { 
            border: 1px solid #444; padding: 12px; text-align: left; 
            vertical-align: top; line-height: 1.7; font-size: 10.5pt;
            white-space: pre-wrap; 
        }
        .analysis-title { font-size: 1.3em; font-weight: bold; margin-top: 30px; margin-bottom: 15px; border-left: 6px solid #000; padding-left: 12px; }

        /* 배경지식 박스 스타일 */
        .background-box { 
            border: 2px solid #2c3e50; 
            padding: 25px; 
            margin-top: 50px; 
            background-color: #f8f9fa; 
            border-radius: 10px; 
            page-break-inside: avoid;
        }
        .background-title { 
            font-size: 1.2em; 
            font-weight: bold; 
            color: #fff; 
            background-color: #2c3e50; 
            padding: 5px 15px; 
            display: inline-block; 
            border-radius: 5px; 
            margin-bottom: 15px; 
        }
        
        .type-box { margin-bottom: 30px; page-break-inside: avoid; }
        h3 { font-size: 1.2em; color: #000; border-bottom: 2px solid #000; padding-bottom: 5px; margin-bottom: 20px; font-weight: bold; margin-top: 40px; } 

        .question-box { margin-bottom: 20px; page-break-inside: avoid; }
        .question-text { font-weight: bold; margin-bottom: 15px; display: block; font-size: 1.1em; word-break: keep-all;} 

        .example-box { 
            border: 1px solid #444; 
            padding: 15px; 
            margin: 15px 0 20px 0; 
            background-color: #fff; 
            font-size: 0.95em; 
            position: relative;
        }
        .example-box::before {
            content: "< 보 기 >";
            display: block;
            text-align: center;
            font-weight: bold;
            color: #333;
            margin-bottom: 10px;
        } 

        .choices { 
            margin-top: 15px; 
            font-size: 1em; 
            margin-left: 15px; 
        }
        .choices div { 
            margin-bottom: 8px; 
            padding-left: 15px; 
            text-indent: -15px; 
            cursor: pointer;
        }
        .choices div:hover { background-color: #f8f9fa; } 

        .write-box { 
            margin-top: 15px; height: 120px; 
            border: 1px solid #ccc; border-radius: 4px;
            background: repeating-linear-gradient(transparent, transparent 29px, #eee 30px); 
            line-height: 30px; 
        } 

        .summary-blank {
            border: 1px dashed #aaa; padding: 15px; margin: 15px 0 25px 0;
            min-height: 100px;
            color: #666; font-size: 0.9em; background-color: #fcfcfc;
            font-weight: bold; display: flex; align-items: flex-start;
        } 

        .blank {
            display: inline-block;
            min-width: 80px; 
            border-bottom: 1.5px solid #000;
            margin: 0 5px;
            height: 1.2em;
            vertical-align: middle;
        } 

        .answer-sheet { 
            background: #f8f9fa; padding: 40px; margin-top: 60px; 
            border-top: 4px double #333; 
            page-break-before: always; 
        }
        .ans-main-title {
            font-size: 1.6em; font-weight: bold; text-align: center; 
            margin-bottom: 40px; padding-bottom: 15px; 
            border-bottom: 3px double #999; color: #333;
        }
        .ans-item { 
            margin-bottom: 50px; 
            border-bottom: 1px dashed #ccc; 
            padding-bottom: 30px; 
        }
        
        .ans-type-badge { 
            display: inline-block; 
            background-color: #555; 
            color: #fff; 
            padding: 4px 12px; 
            border-radius: 15px; 
            font-size: 0.85em; 
            font-weight: bold; 
            margin-bottom: 12px; 
        }
        
        .ans-num { 
            font-weight: bold; 
            color: #d63384; 
            font-size: 1.3em; 
            display: block; 
            margin-bottom: 15px; 
        }
        
        .ans-content-title {
            font-weight: bold;
            color: #2c3e50;
            margin-top: 20px;
            margin-bottom: 8px;
            font-size: 1.05em;
            display: block;
            border-left: 4px solid #2c3e50;
            padding-left: 10px;
        }
        
        .ans-text { 
            display: block; 
            margin-left: 5px; 
            color: #333; 
            line-height: 1.8; 
        }
        
        .ans-wrong-box {
            background-color: #fff;
            border: 1px solid #ddd;
            padding: 15px;
            border-radius: 8px;
            margin-top: 10px;
            color: #555;
        } 

        .summary-ans-box { 
            background-color: #e3f2fd; 
            padding: 25px; 
            margin-bottom: 50px; 
            border-radius: 10px; 
            border: 1px solid #90caf9; 
        }
        .summary-ans-title {
            font-weight: bold; color: #1565c0; font-size: 1.2em; 
            margin-bottom: 15px; display: block; text-align: center;
            border-bottom: 1px solid #90caf9; padding-bottom: 10px;
        }
        
        @media print { body { padding: 0; } }
    </style>
</head>
<body>
""" 
HTML_TAIL = """
</body>
</html>
""" 

# ==========================================
# [헬퍼 함수]
# ==========================================
def get_custom_header_html(main_title, topic_info):
    return f"""
    <div class="header-container">
        <div class="top-row">
            <h1 class="main-title">{main_title}</h1>
            <div class="time-box">소요 시간: &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;</div>
        </div>
        <div class="topic-info">주제: {topic_info}</div>
    </div>
    """ 

def generate_content_with_fallback(prompt, generation_config=None, status_placeholder=None):
    last_exception = None
    for model_name in MODEL_PRIORITY:
        try:
            if status_placeholder:
                status_placeholder.info(f"⚡ 생성 중... (사용 모델: {model_name})")
            if model_name.startswith("gpt") or model_name.startswith("o1"):
                if not openai_client:
                    continue
                response = openai_client.chat.completions.create(
                    model=model_name, 
                    messages=[
                        {"role": "system", "content": "당신은 대한민국 수능 국어 출제 위원장입니다."},
                        {"role": "user", "content": prompt}
                    ],
                    max_completion_tokens=8192 if not generation_config else generation_config.max_output_tokens,
                    temperature=0.7 if not generation_config else generation_config.temperature
                )
                class OpenAIResponseWrapper:
                    def __init__(self, text_content):
                        self.text = text_content
                return OpenAIResponseWrapper(response.choices[0].message.content)
            else:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt, generation_config=generation_config)
                return response
        except Exception as e:
            last_exception = e
            continue 
    if last_exception:
        raise last_exception
    else:
        raise Exception("모델 응답 실패")

def create_docx(html_content, file_name, main_title, topic_title):
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Batang'
    style.font.size = Pt(10) 
    clean_text = re.sub(r'<[^>]+>', '\n', html_content)
    clean_text = re.sub(r'\n+', '\n', clean_text).strip()
    h1 = document.add_heading(main_title, 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_time = document.add_paragraph("소요 시간: ___________")
    p_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_topic = document.add_paragraph(f"주제: {topic_title}")
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("-" * 50)
    document.add_paragraph(clean_text) 
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream 

# ==========================================
# 🧩 1. 비문학 문제 제작 함수 (원본 100% 보존 + 기능 추가)
# ==========================================
def non_fiction_app():
    global GOOGLE_API_KEY
    current_d_mode = st.session_state.get('domain_mode_select', 'AI 생성')
    
    with st.sidebar:
        st.header("🏫 문서 타이틀 설정")
        custom_main_title = st.text_input("메인 타이틀 (학원명)", value="사계국어 모의고사", key="nf_title")
        st.header("🛠️ 지문 입력 방식")
        st.selectbox("방식 선택", ["AI 생성", "직접 입력"], key="domain_mode_select")
        
        # [신규] 지문 출력 여부 선택 기능
        show_passage = st.checkbox("문제지에 지문 포함", value=True, key="nf_show_p")
        # [신규] 배경지식 추가 여부 선택
        use_background = st.checkbox("🔍 지문 관련 배경지식 설명 추가", value=False, key="select_bg_info")
        
        st.header("1️⃣ 지문 및 주제 설정")
        if current_d_mode == 'AI 생성':
            mode = st.radio("구성", ["단일 지문", "주제 통합"], key="ai_mode")
            if mode == "단일 지문":
                domain = st.selectbox("영역", ["인문", "사회", "과학", "기술", "예술"], key="domain_select")
                topic = st.text_input("주제", placeholder="예: 금리 인하", key="topic_input")
                current_topic = topic
                current_domain = domain
            else:
                current_domain = "주제 통합"
                topic_a = st.text_input("주제 (가)", placeholder="예: 공리주의", key="t_a")
                topic_b = st.text_input("주제 (나)", placeholder="예: 의무론", key="t_b")
                current_topic = "(가) " + topic_a + " / (나) " + topic_b
            difficulty = st.select_slider("난이도", ["중", "상", "최상"], value="최상")
            current_difficulty = difficulty
        else: 
            mode = st.radio("지문 구성", ["단일 지문", "주제 통합"], key="manual_mode")
            current_topic = "사용자 지문"
            current_difficulty = "사용자 지정" 

        st.header("2️⃣ 문제 유형 및 개수 선택")
        if mode.startswith("단일"):
            label_type1 = "1. 핵심 주장 요약 (서술형)"
        else:
            label_type1 = "1. (가),(나) 요약 및 연관성 서술"
        
        select_t1 = st.checkbox(label_type1, value=True, key="select_t1")
        select_t2 = st.checkbox("2. 내용 일치 O/X", key="select_t2"); count_t2 = st.number_input(" - 문항 수", 1, 10, 2, key="t2") if select_t2 else 0
        select_t3 = st.checkbox("3. 빈칸 채우기", key="select_t3"); count_t3 = st.number_input(" - 문항 수", 1, 10, 2, key="t3") if select_t3 else 0
        select_t4 = st.checkbox("4. 변형 문장 정오판단", key="select_t4"); count_t4 = st.number_input(" - 문항 수", 1, 10, 2, key="t4") if select_t4 else 0
        select_t5 = st.checkbox("5. 객관식 (일치/불일치)", value=True, key="select_t5"); count_t5 = st.number_input(" - 문항 수", 1, 10, 2, key="t5") if select_t5 else 0
        select_t6 = st.checkbox("6. 객관식 (추론)", value=True, key="select_t6"); count_t6 = st.number_input(" - 문항 수", 1, 10, 2, key="t6") if select_t6 else 0
        select_t7 = st.checkbox("7. 객관식 (보기 적용 3점)", value=True, key="select_t7"); count_t7 = st.number_input(" - 문항 수", 1, 10, 1, key="t7") if select_t7 else 0
        
        use_summary = st.checkbox("📌 문단별 요약 훈련 칸 생성", value=True, key="select_summary")

    if st.session_state.generation_requested:
        manual_p = ""
        if current_d_mode == '직접 입력':
            if mode == '단일 지문':
                manual_p = st.session_state.get("manual_passage_input_col_main", "")
            else:
                manual_p = "[가] 지문:\n" + st.session_state.get("manual_passage_input_a", "") + "\n\n[나] 지문:\n" + st.session_state.get("manual_passage_input_b", "")

        if not current_topic and current_d_mode == 'AI 생성':
            st.warning("주제를 입력해주세요."); st.session_state.generation_requested = False; return
        elif current_d_mode == '직접 입력' and not manual_p.strip():
            st.warning("지문을 입력해주세요."); st.session_state.generation_requested = False; return
        else:
            status = st.empty(); status.info(f"⚡ 출제 준비 중...")
            try:
                # [복구] 상세 문항 가이드라인
                req_list = []
                if select_t1: req_list.append('<div class="question-box"><span class="question-text">1. ' + label_type1 + '</span><div class="write-box"></div></div><br><br>')
                if select_t2: req_list.append('<h3>내용 일치 O/X (' + str(count_t2) + '문항)</h3>- 문항 끝에 ( O / X ) 포함. 각 문제 뒤에 <br><br> 삽입.')
                if select_t3: req_list.append('<h3>빈칸 채우기 (' + str(count_t3) + '문항)</h3>- 빈칸은 `<span class="blank">&nbsp;&nbsp;&nbsp;&nbsp;</span>` 사용. 각 문제 뒤에 <br><br> 삽입.')
                if select_t4: req_list.append('<h3>변형 문장 정오판단 (' + str(count_t4) + '문항)</h3>- 문항 끝에 ( O / X ) 포함. 각 문제 뒤에 <br><br> 삽입.')
                mcq_tpl = '<div class="question-box"><span class="question-text">[문제번호] [발문]</span><div class="choices"><div>① ...</div><div>② ...</div><div>③ ...</div><div>④ ...</div><div>⑤ ...</div></div></div><br><br>'
                if select_t5: req_list.append('<h3>객관식: 세부 내용 파악 (' + str(count_t5) + '문항)</h3>' + mcq_tpl)
                if select_t6: req_list.append('<h3>객관식: 추론 및 비판 (' + str(count_t6) + '문항)</h3>' + mcq_tpl)
                if select_t7: req_list.append('<h3>객관식: [보기] 적용 문제 (' + str(count_t7) + '문항) [3점]</h3><div class="question-box"><span class="question-text">[문제번호] 윗글을 바탕으로 [보기]를 이해한 내용으로 적절하지 않은 것은? [3점]</span><div class="example-box">(보기 내용)</div><div class="choices"><div>① ...</div><div>② ...</div><div>③ ...</div><div>④ ...</div><div>⑤ ...</div></div></div><br><br>')
                
                reqs_str = "\n".join(req_list)
                
                # [복구] 문단 요약 및 배경지식 프롬프트
                summary_inst_passage = """
                - **[필수]**: 각 문단이 끝날 때마다 반드시 `<div class='summary-blank'>📝 문단 요약 연습: (이곳에 핵심 내용을 요약해보세요)</div>` 코드를 삽입하여 사용자가 내용을 요약할 수 있는 빈칸을 만들어주시오.
                - 이 부분은 사용자가 글을 쓸 공간이므로 절대 내용을 채우지 마시오.
                """ if use_summary else ""

                bg_instruction = ""
                if use_background:
                    bg_instruction = """
                    - **[배경지식 플러스 서술]**: 모든 문제 출제가 끝난 후, 맨 마지막에 지문의 주제와 관련된 심화 배경지식을 정리하시오.
                    - 제목은 `<div class="background-title">💡 배경지식 플러스</div>`로 하고, 전체 내용은 `<div class="background-box">`로 감싸시오.
                    - 지문에서 다룬 원리나 사건의 유래, 현실 세계의 적용 사례 등을 500자 내외로 상세하고 친절하게 설명하시오.
                    """

                # [원본 유지] 지문 가이드라인 및 킬러 가이드
                # [추가] 직접 입력 모드에서 지문을 중복 출력하지 않도록 명시
                p1_prompt = """
당신은 대한민국 수능 국어 출제 위원장입니다. 
아래 지시사항에 맞춰 완벽한 HTML 포맷의 모의고사 문제지를 생성하시오.
- `<html>`, `<head>` 생략, `<body>` 내용만 출력.
- 정답 및 해설 제외. 학생용 문제지.
# 🚨 [매우 중요] 출력 시 절대 제목/헤더를 생성하지 마시오. h1, h2 태그 및 제목 노출 금지.

{STEP1}
{USER_BLOCK}
{BG_PROM}

# ----------------------------------------------------------------
# 🚨 [고난도(킬러 문항) 출제 필수 가이드라인]
# ----------------------------------------------------------------
1. **[정보의 재구성 필수 - 1:1 매칭 금지]**: 
   - 정답 선지는 절대 한 문단이나 한 문장의 내용만으로 판단할 수 없게 하시오.
   - **반드시 '1문단 + 3문단' 혹은 'A주장 + B반론'처럼 서로 멀리 떨어진 두 개 이상의 정보를 결합**해야만 참/거짓을 판별할 수 있도록 문장을 재구성하시오.

2. **[단어 바꿔치기(Paraphrasing)]**:
   - 지문에 있는 단어를 그대로 선지에 쓰지 마시오.
   - 지문의 '상승했다'를 '하락하지 않았다'나 '고점에 도달했다'처럼 **동의어나 함축적 의미로 변환**하여 선지를 작성하시오.

3. **[인과관계 비틀기 (오답 설계)]**:
   - 단순히 '아니다'를 붙이는 유치한 오답을 금지합니다.
   - 'A라서 B이다'를 'B라서 A이다'로 **인과관계를 뒤집거나**, 주체(주어)와 객체(목적어)를 서로 바꾸어 매력적인 오답을 만드시오.

4. **[선지 분포]**:
   - 선지 ①~⑤번이 지문의 특정 부분에 쏠리지 않게, 지문 전체(서론, 본론, 결론)를 아우르도록 배치하시오.

**[Step 2] 문제 출제**
{REQS}
                """.format(
                    STEP1 = f"**[Step 1] 지문 작성** - 주제: {current_topic}, 난이도: {difficulty}, 길이: 1800자 내외. 생성된 지문은 반드시 `<div class='passage'>` 태그로 감싸시오. \n{summary_inst_passage}" if current_d_mode == 'AI 생성' else "**[Step 1] 지문 인식** - 사용자 입력 지문 기반. 문제지 본문에는 지문을 다시 출력하지 마시오.",
                    USER_BLOCK = "\n[사용자 입력 지문 시작]\n" + manual_p + "\n[사용자 입력 지문 끝]\n" if current_d_mode == '직접 입력' else "",
                    BG_PROM = bg_instruction,
                    REQS = reqs_str
                )
                
                res_q = generate_content_with_fallback(p1_prompt, status_placeholder=status)
                html_q = res_q.text.replace("```html", "").replace("```", "").strip()
                html_q = re.sub(r'<h[12].*?>.*?</h[12]>', '', html_q, flags=re.DOTALL | re.IGNORECASE)

                # [복구] 해설 분할 생성 (Batch Size 6) 로직
                total_q_cnt = sum([1 if select_t1 else 0, count_t2, count_t3, count_t4, count_t5, count_t6, count_t7])
                BATCH_SIZE = 6; final_ans_parts = []; summary_done = False
                extra_context = "\n**[참고: 지문 원문]**\n" + manual_p + "\n" if current_d_mode == '직접 입력' else ""

                for i in range(0, total_q_cnt, BATCH_SIZE):
                    start_num = i + 1; end_num = min(i + BATCH_SIZE, total_q_cnt)
                    status.info(f"📝 정답 및 해설 생성 중... ({start_num}~{end_num}번 / 총 {total_q_cnt}문항)")
                    
                    current_summary_prompt = ""
                    if use_summary and not summary_done:
                        # [수정] AI에게 구조적 개조식 요약을 강제하는 상세 지침 (원본 유지)
                        structure_inst = (
                            "단순한 서술형 문장이 아니라, 정보를 명확히 분류한 **[개조식]** 형태로 요약하시오. "
                            "반드시 **1. 핵심 화제, 2. 논리적 전개 방식(정의, 대조, 인과 등), 3. 핵심 요지** 항목을 명확히 구분하여 "
                            "'- 화제: [내용] / - 방식: [내용] / - 요지: [내용]'과 같이 구조화된 형식을 사용하여 가독성을 높이시오."
                        )
                        if current_d_mode == '직접 입력':
                             p_cnt = len([p for p in re.split(r'\n\s*\n', manual_p.strip()) if p.strip()])
                             current_summary_prompt = f"- **[필수 - 최우선 작성]**: 답변 맨 위에 `<div class='summary-ans-box'>`를 열고 **[문단별 구조적 요약 예시 답안]**을 작성하시오. 총 {p_cnt}개의 문단 요약을 제시하시오. 지침: {structure_inst}"
                        else:
                             current_summary_prompt = f"- **[필수 - 최우선 작성]**: 답변 맨 위에 `<div class='summary-ans-box'>`를 열고 **[문단별 구조적 요약 예시 답안]**을 작성하시오. 지침: {structure_inst}"
                        summary_done = True 

                    p_chunk = """
당신은 대한민국 수능 국어 출제 위원장입니다. {T_CNT}문제 중 **{S_NUM}번부터 {E_NUM}번까지**의 정답 및 해설을 HTML로 작성하시오.
{CONTEXT}
[입력된 문제]: {Q_TEXT}
{SUM_PROM}
[규칙]: 객관식은 정답 상세 해설 + 오답 분석 필수. OX/빈칸은 지문 근거 필수.
                    """.format(T_CNT=total_q_cnt, S_NUM=start_num, E_NUM=end_num, CONTEXT=extra_context, Q_TEXT=html_q, SUM_PROM=current_summary_prompt)
                    
                    res_chunk = generate_content_with_fallback(p_chunk, status_placeholder=status)
                    chunk_text = res_chunk.text.replace("```html","").replace("```","").strip()
                    if i == 0: chunk_text = '<div class="answer-sheet"><h2 class="ans-main-title">정답 및 해설</h2>' + chunk_text
                    final_ans_parts.append(chunk_text)

                html_answers = "".join(final_ans_parts) + "</div>"
                full_html = HTML_HEAD + get_custom_header_html(custom_main_title, current_topic)
                
                # [지문 출력 제어 강화 로직]
                processed_html_q = html_q
                if not show_passage:
                    # AI가 생성한 문자열 내부에서 passage 클래스를 가진 영역을 정규식으로 완벽 제거
                    processed_html_q = re.sub(r'<div[^>]*class=["\']passage["\'][^>]*>.*?</div>', '', html_q, flags=re.DOTALL | re.IGNORECASE)
                    full_html += processed_html_q
                else:
                    if current_d_mode == '직접 입력':
                        paras = [p.strip() for p in re.split(r'\n\s*\n', manual_p.strip()) if p.strip()]
                        formatted_p = "".join([f"<p>{p}</p>" + ("<div class='summary-blank'>📝 문단 요약 연습: </div>" if use_summary else "") for p in paras])
                        full_html += f'<div class="passage">{formatted_p}</div>' + processed_html_q
                    else:
                        # AI 생성 모드에서는 AI가 이미 <div class='passage'>를 생성했으므로 그대로 출력
                        full_html += processed_html_q
                
                full_html += html_answers + HTML_TAIL
                st.session_state.generated_result = {"full_html": full_html, "main_title": custom_main_title, "topic_title": current_topic}
                status.success("✅ 비문학 생성 완료!"); st.session_state.generation_requested = False
            except Exception as e: status.error(f"오류: {e}"); st.session_state.generation_requested = False

# ==========================================
# 📖 2. 소설 문제 제작 함수 (원본 100% 보존 + 기능 추가)
# ==========================================
def fiction_app():
    with st.sidebar:
        st.header("🏫 문서 타이틀 설정")
        custom_main_title = st.text_input("메인 타이틀 (학원명)", value="사계국어 모의고사", key="fic_t")
        st.header("🛠️ 출력 설정")
        show_passage = st.checkbox("문제지에 지문 포함", value=True, key="fi_show_p")
        st.header("1️⃣ 작품 정보"); work_name = st.text_input("작품명", key="fic_n"); author_name = st.text_input("작가명", key="fic_a")
        st.header("2️⃣ 문제 유형 및 개수")
        uv = st.checkbox("1. 어휘 문제 (단답형)", value=True, key="fv"); cv = st.number_input("문항 수", 1, 20, 5, key="fcv") if uv else 0
        ue = st.checkbox("2. 서술형 심화 (감상)", value=True, key="fe"); ce = st.number_input("문항 수", 1, 10, 3, key="fce") if ue else 0
        um = st.checkbox("3. 객관식 (일반)", value=True, key="fm"); cm = st.number_input("문항 수", 1, 10, 3, key="fcm") if um else 0
        ub = st.checkbox("4. 객관식 (보기 적용)", value=True, key="fb"); cb = st.number_input("문항 수", 1, 10, 2, key="fcb") if ub else 0
        st.caption("3️⃣ 분석 및 정리 활동 (서술형/표)")
        u5 = st.checkbox("5. 주요 등장인물 정리 (표)", key="f5"); u6 = st.checkbox("6. 소설 속 상황 요약", key="f6")
        u7 = st.checkbox("7. 인물 관계도 및 갈등", key="f7"); u8 = st.checkbox("8. 갈등 구조 및 심리 정리", key="f8")

    if st.session_state.generation_requested:
        text = st.session_state.fiction_novel_text_input_area
        if not text: st.warning("본문을 입력하세요."); st.session_state.generation_requested = False; return
        status = st.empty(); status.info("⚡ 소설 심층 분석 및 문제 제작 중...")
        try:
            req_list = []
            if uv: req_list.append('<div class="type-box"><h3>유형 1. 어휘 문제 (' + str(cv) + '문항)</h3>- 지문의 어려운 어휘 ' + str(cv) + '개의 의미 묻기 (단답형).<div class="question-box"><span class="question-text">[번호] "____"의 문맥적 의미는?</span><div class="write-box" style="height:50px;"></div></div></div><br><br>')
            if ue: req_list.append('<div class="type-box"><h3>유형 2. 서술형 심화 문제 (' + str(ce) + '문항)</h3>- 작가의 의도, 효과, 이유를 묻는 고난도 서술형.<div class="write-box"></div></div><br><br>')
            if um: req_list.append('<div class="type-box"><h3>유형 3. 객관식 문제 (일반) (' + str(cm) + '문항)</h3>- 수능형 5지 선다 (추론/비판).<div class="question-box"><span class="question-text">[번호] (발문)</span><div class="choices"><div>① ...</div><div>② ...</div><div>③ ...</div><div>④ ...</div><div>⑤ ...</div></div></div></div><br><br>')
            if ub: req_list.append('<div class="type-box"><h3>유형 4. 객관식 문제 (보기 적용) (' + str(cb) + '문항)</h3>- **<보기>** 박스 필수 포함 (3점 킬러문항).<div class="example-box">(보기 내용)</div><div class="choices"><div>① ...</div><div>② ...</div><div>③ ...</div><div>④ ...</div><div>⑤ ...</div></div></div></div><br><br>')
            if u5: req_list.append('<div class="type-box"><h3>유형 5. 주요 등장인물 정리</h3>- 인물명, 호칭, 역할, 심리 빈칸 표 제공.</div><br><br>')
            if u6: req_list.append('<div class="type-box"><h3>유형 6. 소설 속 상황 요약</h3>- 핵심 갈등 요약 서술.<div class="write-box"></div></div><br><br>')
            if u7: req_list.append('<div class="type-box"><h3>유형 7. 인물 관계도 및 갈등</h3>- 직접 그릴 수 있는 박스.<div class="write-box" style="height:200px;"></div></div><br><br>')
            if u8: req_list.append('<div class="type-box"><h3>유형 8. 갈등 구조 및 심리 정리</h3>- 갈등 양상 및 비판 의도 서술.<div class="write-box"></div></div><br><br>')
            
            r_str = "\n".join(req_list)
            p1_p = """
당신은 수능 문학 출제위원입니다. 작품 '{W_N}'({A_N}) 기반 학생용 문제지(HTML)를 작성하시오.
# 🚨 [수능 최고난도 출제 지침]
1. **[복합적 사고]**: 작품 전체 맥락과 함축적 의미를 종합해야 풀 수 있는 문제.
2. **[매력적인 오답]**: 부분적 진실, 주객 전도, 과잉 해석 함정 배치.
3. **[보기 적용]**: 비평적 관점을 적용해 새롭게 해석하는 3점 문항.
4. **[가독성 개선]**: 모든 문항 뒤에 <br><br>을 삽입하시오.

# 🚨 [매우 중요] h1, h2 제목 생성 금지. 본문 내용부터 바로 출력. 지문 본문은 절대 포함하지 마시오.
본문: {BODY}
[출제 요청 목록]:
{REQS}
            """.format(W_N=work_name, A_N=author_name, BODY=text, REQS=r_str)
            
            res_q = generate_content_with_fallback(p1_p, status_placeholder=status)
            html_q = res_q.text.replace("```html","").replace("```","").strip()
            html_q = re.sub(r'<h[12].*?>.*?</h[12]>', '', html_q, flags=re.DOTALL | re.IGNORECASE)

            p2_p = """
당신은 수능 문학 해설 위원입니다. 앞서 출제된 문제들에 대한 **완벽한 정답 및 해설**을 <div class="answer-sheet"> 내부에 작성하시오.
**[작성 규칙]**: 1. 객관식은 [정답], [상세 해설], [오답 분석] 필수. 2. 활동형은 예시 답안 제시.
[입력 문제 내용]: {Q_TEXT}
            """.format(Q_TEXT=html_q)
            res_a = generate_content_with_fallback(p2_p, status_placeholder=status)
            html_a = res_a.text.replace("```html","").replace("```","").strip()
            
            full_html = HTML_HEAD + get_custom_header_html(custom_main_title, work_name)
            
            # [지문 출력 완벽 제어]
            processed_html_q = html_q
            if not show_passage:
                processed_html_q = re.sub(r'<div[^>]*class=["\']passage["\'][^>]*>.*?</div>', '', html_q, flags=re.DOTALL | re.IGNORECASE)
                full_html += processed_html_q
            else:
                full_html += f'<div class="passage">{text.replace(chr(10), "<br>")}</div>' + processed_html_q
            
            full_html += html_a + HTML_TAIL
            st.session_state.generated_result = {"full_html": full_html, "main_title": custom_main_title, "topic_title": work_name}
            status.success("✅ 소설 분석 완료!"); st.session_state.generation_requested = False
        except Exception as e: status.error(f"오류: {e}"); st.session_state.generation_requested = False

# ==========================================
# 🌸 3. 운문 분석 차트형 분석 및 고난도 문항 제작
# ==========================================
def poetry_app():
    with st.sidebar:
        st.header("🏫 문서 타이틀 설정")
        c_title = st.text_input("메인 타이틀", value="사계국어 모의고사", key="po_t")
        st.header("🛠️ 출력 설정")
        show_passage = st.checkbox("문제지에 지문 포함", value=True, key="po_show_p")
        st.header("1️⃣ 작품 정보"); po_n = st.text_input("작품명", key="po_n"); po_a = st.text_input("작가명", key="po_a")
        
        # [원본 요청] 갈래 선택 기능
        po_genre = st.selectbox("작품 갈래", ["현대시", "고대가요", "향가", "고려가요", "시조", "가사", "악장", "잡가", "민요", "한시"], key="po_g")
        
        st.header("2️⃣ 분석 차트 구성 (1~7번 자동생성)")
        st.caption("개요~키포인트 자동생성")
        # [원본 요청] 어휘 풀이 선택 옵션
        ct_vocab_analysis = st.checkbox("7. 주요 어휘 및 구절 풀이", value=True, key="po_vocab_opt")
        
        st.header("3️⃣ 문제 제작 및 개수")
        ct8 = st.checkbox("8. 수능형 선지 O,X 세트", value=True); nt8 = st.number_input("OX 문항 수", 1, 15, 10, key="pn8") if ct8 else 0
        ct9 = st.checkbox("9. 수능형 서술형 문제", value=True); nt9 = st.number_input("서술형 문항 수", 1, 10, 3, key="pn9") if ct9 else 0

    if st.session_state.generation_requested:
        text = st.session_state.get("poetry_text_input_area", "")
        if not text: st.warning("운문 본문을 입력하세요."); st.session_state.generation_requested = False; return
        status = st.empty(); status.info("⚡ 운문 분석 중...")
        try:
            # [복구] 어휘 풀이 행 동적 생성
            vocab_row = ""
            if ct_vocab_analysis:
                vocab_row = "  <tr><th>7. 주요 어휘 및 구절 풀이</th><td>(지문 속 중요 어휘나 난해한 구절을 상세히 풀이)</td></tr>"

            # [원본 유지] 분석 차트 프롬프트
            p_chart = """
당신은 수능 국어 강사입니다. 운문 작품 '{W_N}'({A_N}, 갈래: {G_N})를 분석하여 아래 HTML 차트를 제작하시오.
[포맷 지침]: 반드시 아래 HTML 구조를 엄격히 지켜서 출력할 것.
1. 사용자가 설정한 갈래인 '{G_N}'의 특성을 정확히 반영하여 분석하시오.
2. 각 항목의 내용은 1), 2), 3) 과 같은 순서 표시를 사용하여 요점 위주로 작성하시오.
3. 내용이 길어질 경우 적절한 줄바꿈을 포함하여 가독성을 높이시오.
4. 제목 칸(th)의 너비는 120px로 고정되도록 디자인 지침을 따르시오.

<div class="analysis-title">운문 분석 : {W_N} ({G_N})</div>
<table class="analysis-chart">
  <tr><th>1. 작품 개요</th><td>(갈래 {G_N}의 형식적 특징, 성격, 주제 등을 상세히 기술)</td></tr>
  <tr><th>2. 핵심 내용 정리</th><td>(시상 전개 과정 및 핵심 상황 요약)</td></tr>
  <tr><th>3. 주요 소재의 상징성</th><td>(주요 시어 및 비유적 소재의 의미 분석)</td></tr>
  <tr><th>4. 표현상의 특징</th><td>(사용된 수사법, 심상, 어조, {G_N} 특유의 율격 특징)</td></tr>
  <tr><th>5. 작품의 이해와 감상</th><td>(작품의 문학적 가치와 종합적 감상평)</td></tr>
  <tr><th>6. 수능의 키포인트</th><td>(이 작품에서 수능 고난도 킬러 문항으로 출제될 수 있는 포인트)</td></tr>
{V_ROW}
</table>
본문: {BODY}
            """.format(W_N=po_n, A_N=po_a, G_N=po_genre, BODY=text, V_ROW=vocab_row)
            
            res_chart = generate_content_with_fallback(p_chart, status_placeholder=status)
            html_chart = res_chart.text.replace("```html","").replace("```","").strip()

            # [원본 유지] 문제 생성 프롬프트
            r_list = []
            if ct8: r_list.append("문항 8. 수능형 선지 OX 판단 (" + str(nt8) + "개) - 질문 끝에 ( ) 빈칸 출력. 각 문항 뒤 <br><br> 필수.")
            if ct9: r_list.append("문항 9. 고난도 수능형 서술형 (" + str(nt9) + "개) - 각 문항 뒤 <br><br> 필수.")
            r_str = "\n".join(r_list)
            
            p_q = """
당신은 수능 국어 출제 위원장입니다. 운문 작품 '{W_N}'(갈래: {G_N})를 바탕으로 학생용 문제지(HTML)를 제작하시오.
[중요 지침]: 
1. {G_N}의 장르적 특성을 고려하여 실제 수능형 문제를 출제하시오. 
2. 출력 시 반드시 아래의 HTML 구조를 따를 것: 각 문항은 question-box 클래스를 사용하고 문항 끝에는 <br><br>을 삽입하시오.
3. 시 본문은 이미 출력했으므로 **HTML 응답에 절대 시 본문을 포함하지 마시오.** 출제 요청:
{REQS}
본문: {BODY}
            """.format(W_N=po_n, G_N=po_genre, REQS=r_str, BODY=text)
            
            res_q = generate_content_with_fallback(p_q, status_placeholder=status)
            html_q = res_q.text.replace("```html","").replace("```","").strip()
            html_q = re.sub(r'<h[12].*?>.*?</h[12]>', '', html_q, flags=re.DOTALL | re.IGNORECASE)

            p_a = "위 8~9번 문항들에 대해 교사용 완벽 정답 및 상세 해설을 <div class='answer-sheet'> 내부에 작성하시오.\n문제 내용: " + html_q
            res_a = generate_content_with_fallback(p_a, status_placeholder=status)
            html_a = res_a.text.replace("```html","").replace("```","").strip()
            
            full_html = HTML_HEAD + get_custom_header_html(c_title, po_n)
            
            # [지문 출력 완벽 제어]
            processed_html_q = html_q
            if not show_passage:
                processed_html_q = re.sub(r'<div[^>]*class=["\']poetry-passage["\'][^>]*>.*?</div>', '', html_q, flags=re.DOTALL | re.IGNORECASE)
                full_html += processed_html_q
            else:
                full_html += f'<div class="poetry-passage">{text}</div>' + processed_html_q
            
            full_html += html_chart + html_a + HTML_TAIL
            st.session_state.generated_result = {"full_html": full_html, "main_title": c_title, "topic_title": po_n}
            status.success("✅ 운문 분석 완료!"); st.session_state.generation_requested = False
        except Exception as e: status.error(f"오류: {e}"); st.session_state.generation_requested = False

# ==========================================
# 🚀 메인 실행 로직
# ==========================================
def display_results():
    if st.session_state.generated_result:
        res = st.session_state.generated_result
        st.markdown("---")
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("🔄 다시 생성"):
                st.session_state.generated_result = None; st.session_state.generation_requested = True; st.rerun()
        with c2: st.download_button("📥 HTML 저장", res["full_html"], "exam.html", "text/html")
        with c3:
            docx = create_docx(res["full_html"], "exam.docx", res["main_title"], res["topic_title"])
            st.download_button("📄 Word 저장", docx, "exam.docx")
        st.components.v1.html(res["full_html"], height=800, scrolling=True)

st.title("📚 사계국어 모의고사 제작 시스템")
st.markdown("---")
col_L, col_R = st.columns([1.5, 3])

with col_L:
    st.radio("모드 선택", ["⚡ 비문학 문제 제작", "📖 소설 문제 제작", "🌸 운문 문제 제작"], key="app_mode")

with col_R:
    if st.session_state.app_mode == "⚡ 비문학 문제 제작":
        st.header("⚡ 비문학 모의평가")
        if st.session_state.get("domain_mode_select") == "직접 입력":
            m_m = st.session_state.get("manual_mode", "단일 지문")
            if m_m == "단일 지문": st.text_area("지문 입력 (엔터 두번으로 문단 구분)", height=300, key="manual_passage_input_col_main")
            else:
                ca, cb = st.columns(2)
                with ca: st.text_area("(가) 지문", height=300, key="manual_passage_input_a")
                with cb: st.text_area("(나) 지문", height=300, key="manual_passage_input_b")
        if st.button("🚀 모의고사 생성", key="r_nf"): st.session_state.generation_requested = True
        non_fiction_app()
    elif st.session_state.app_mode == "🌸 운문 문제 제작":
        st.header("🌸 운문 분석 차트 및 문항 제작")
        st.text_area("운문 본문 입력 (행/연 구분 정확히)", height=400, key="poetry_text_input_area")
        if st.button("🚀 분석 및 제작 시작", key="r_po"): st.session_state.generation_requested = True
        poetry_app()
    else:
        st.header("📖 소설 심층 분석")
        st.text_area("작품 본문 입력", height=300, key="fiction_novel_text_input_area")
        if st.button("🚀 분석 생성", key="r_fi"): st.session_state.generation_requested = True
        fiction_app()

display_results()
